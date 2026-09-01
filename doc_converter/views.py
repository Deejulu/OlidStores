import os
import uuid
import time
import re
import json
import logging
from pathlib import Path

from django.conf import settings
from django.core.cache import cache
from django.shortcuts import render, redirect
from django.http import JsonResponse, HttpResponse, Http404
from django.utils import timezone
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt

logger = logging.getLogger('doc_converter')

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
CACHE_TIMEOUT = 3600  # 1 hour
ALLOWED_EXTENSIONS = {'pdf', 'docx'}


def _get_upload_dir():
    return os.path.join(settings.MEDIA_ROOT, 'doc_converter', 'uploads')


def _get_output_dir():
    return os.path.join(settings.MEDIA_ROOT, 'doc_converter', 'outputs')


def _ensure_dirs():
    os.makedirs(_get_upload_dir(), exist_ok=True)
    os.makedirs(_get_output_dir(), exist_ok=True)


def _cleanup_old_files(directory, max_age_seconds=3600):
    now = time.time()
    if not os.path.exists(directory):
        return
    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        if os.path.isfile(filepath):
            try:
                mtime = os.path.getmtime(filepath)
                if now - mtime > max_age_seconds:
                    os.remove(filepath)
            except OSError:
                pass


def _cleanup_task(task_id):
    for ext in ('pdf', 'docx'):
        upload_path = os.path.join(_get_upload_dir(), f'{task_id}.{ext}')
        output_path = os.path.join(_get_output_dir(), f'{task_id}.{ext}')
        for p in (upload_path, output_path):
            try:
                if os.path.exists(p):
                    os.remove(p)
            except OSError:
                pass


def _detect_file_type(filename):
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    if ext not in ALLOWED_EXTENSIONS:
        return None
    return ext


def _convert_pdf_to_docx(upload_path, output_path):
    from pdf2docx import Converter
    cv = Converter(upload_path)
    cv.convert(output_path)
    cv.close()


def _convert_docx_to_pdf(upload_path, output_path):
    from docx2pdf import convert
    convert(upload_path, output_path)


def _extract_docx_text(upload_path):
    from docx import Document
    doc = Document(upload_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else 'Normal'
            if 'Heading' in style:
                level = style.replace('Heading ', '').replace('Heading', '1')
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                paragraphs.append({'type': 'heading', 'level': level, 'text': text})
            elif 'List' in style:
                paragraphs.append({'type': 'list', 'text': text})
            else:
                paragraphs.append({'type': 'paragraph', 'text': text})
    return paragraphs


def _create_docx_from_text(paragraphs, output_path):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    for item in paragraphs:
        text = item.get('text', '')
        if not text:
            continue
        ptype = item.get('type', 'paragraph')
        if ptype == 'heading':
            level = item.get('level', 1)
            doc.add_heading(text, level=min(level, 9))
        elif ptype == 'list':
            doc.add_paragraph(text, style='List Bullet')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    doc.save(output_path)


def index(request):
    return render(request, 'doc_converter/index.html')


@require_POST
@csrf_exempt
def upload_file(request):
    _ensure_dirs()
    _cleanup_old_files(_get_upload_dir())
    _cleanup_old_files(_get_output_dir())

    uploaded_file = request.FILES.get('file')
    if not uploaded_file:
        return JsonResponse({'success': False, 'error': 'No file was uploaded.'}, status=400)

    ext = _detect_file_type(uploaded_file.name)
    if not ext:
        return JsonResponse({'success': False, 'error': 'Only PDF and Word (.docx) files are allowed.'}, status=400)

    if uploaded_file.size > MAX_FILE_SIZE:
        max_mb = MAX_FILE_SIZE / (1024 * 1024)
        return JsonResponse({'success': False, 'error': f'File is too large. Maximum size is {max_mb:.0f}MB.'}, status=400)

    task_id = str(uuid.uuid4())
    upload_path = os.path.join(_get_upload_dir(), f'{task_id}.{ext}')

    try:
        with open(upload_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
    except Exception as e:
        logger.exception('Failed to save uploaded file')
        return JsonResponse({'success': False, 'error': 'Failed to save uploaded file. Please try again.'}, status=500)

    cache.set(f'doc_converter_task_{task_id}', {
        'status': 'processing',
        'created_at': timezone.now().isoformat(),
        'original_name': uploaded_file.name,
        'file_type': ext,
    }, timeout=CACHE_TIMEOUT)

    try:
        if ext == 'pdf':
            output_path = os.path.join(_get_output_dir(), f'{task_id}.docx')
            _convert_pdf_to_docx(upload_path, output_path)
            download_url = f'/tools/pdf-to-word/download/{task_id}/'
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            download_name = f'{uploaded_file.name.rsplit(".", 1)[0]}.docx'
        else:
            output_path = os.path.join(_get_output_dir(), f'{task_id}.pdf')
            _convert_docx_to_pdf(upload_path, output_path)
            download_url = f'/tools/pdf-to-word/download/{task_id}/'
            mime = 'application/pdf'
            download_name = f'{uploaded_file.name.rsplit(".", 1)[0]}.pdf'

        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise RuntimeError('Conversion produced an empty file.')

        cache.set(f'doc_converter_task_{task_id}', {
            'status': 'completed',
            'created_at': timezone.now().isoformat(),
            'original_name': uploaded_file.name,
            'file_type': ext,
            'download_url': download_url,
            'mime': mime,
            'download_name': download_name,
        }, timeout=CACHE_TIMEOUT)

        return JsonResponse({
            'success': True,
            'task_id': task_id,
            'file_type': ext,
            'download_url': download_url,
        })

    except Exception as e:
        logger.exception('Conversion failed for task %s', task_id)
        cache.set(f'doc_converter_task_{task_id}', {
            'status': 'failed',
            'created_at': timezone.now().isoformat(),
            'original_name': uploaded_file.name,
            'file_type': ext,
            'error': str(e),
        }, timeout=CACHE_TIMEOUT)
        _cleanup_task(task_id)
        return JsonResponse({'success': False, 'error': 'Conversion failed. The file may be corrupted or password-protected. Please try another file.'}, status=400)


@require_POST
@csrf_exempt
def edit_docx(request, task_id):
    cache_key = f'doc_converter_task_{task_id}'
    task = cache.get(cache_key)
    if not task or task.get('status') != 'completed' or task.get('file_type') != 'docx':
        return JsonResponse({'success': False, 'error': 'Document not found or not available for editing.'}, status=404)

    upload_path = os.path.join(_get_upload_dir(), f'{task_id}.docx')
    if not os.path.exists(upload_path):
        return JsonResponse({'success': False, 'error': 'Original file not found.'}, status=404)

    try:
        paragraphs = _extract_docx_text(upload_path)
        text_lines = []
        for item in paragraphs:
            if item['type'] == 'heading':
                text_lines.append(f'[H{item["level"]}] {item["text"]}')
            elif item['type'] == 'list':
                text_lines.append(f'[LIST] {item["text"]}')
            else:
                text_lines.append(item['text'])
        return JsonResponse({'success': True, 'text': '\n'.join(text_lines)})
    except Exception as e:
        logger.exception('Failed to extract text for task %s', task_id)
        return JsonResponse({'success': False, 'error': 'Failed to read document text.'}, status=500)


@require_POST
@csrf_exempt
def save_edited_docx(request):
    try:
        body = json.loads(request.body.decode('utf-8'))
        task_id = body.get('task_id')
        text = body.get('text', '')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Invalid request.'}, status=400)

    if not task_id:
        return JsonResponse({'success': False, 'error': 'Missing task ID.'}, status=400)

    cache_key = f'doc_converter_task_{task_id}'
    task = cache.get(cache_key)
    if not task or task.get('status') != 'completed' or task.get('file_type') != 'docx':
        return JsonResponse({'success': False, 'error': 'Document not found.'}, status=404)

    if not text.strip():
        return JsonResponse({'success': False, 'error': 'Document text is empty.'}, status=400)

    paragraphs = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        m = re.match(r'\[H(\d)\]\s*(.*)', line)
        if m:
            paragraphs.append({'type': 'heading', 'level': int(m.group(1)), 'text': m.group(2)})
            continue
        m = re.match(r'\[LIST\]\s*(.*)', line)
        if m:
            paragraphs.append({'type': 'list', 'text': m.group(1)})
            continue
        paragraphs.append({'type': 'paragraph', 'text': line})

    new_task_id = str(uuid.uuid4())
    output_path = os.path.join(_get_output_dir(), f'{new_task_id}.docx')
    try:
        _create_docx_from_text(paragraphs, output_path)
    except Exception as e:
        logger.exception('Failed to create edited DOCX for task %s', task_id)
        return JsonResponse({'success': False, 'error': 'Failed to create document. Please try again.'}, status=500)

    original_name = task.get('original_name', 'document')
    base_name = original_name.rsplit('.', 1)[0] if '.' in original_name else original_name
    download_name = f'{base_name}_edited.docx'
    download_url = f'/tools/pdf-to-word/download/{new_task_id}/'

    cache.set(f'doc_converter_task_{new_task_id}', {
        'status': 'completed',
        'created_at': timezone.now().isoformat(),
        'original_name': download_name,
        'file_type': 'docx',
        'download_url': download_url,
        'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'download_name': download_name,
    }, timeout=CACHE_TIMEOUT)

    return JsonResponse({'success': True, 'download_url': download_url})


def download_file(request, task_id):
    cache_key = f'doc_converter_task_{task_id}'
    task = cache.get(cache_key)
    if not task or task.get('status') != 'completed':
        raise Http404('Conversion not found or expired.')

    file_type = task.get('file_type', 'docx')
    download_name = task.get('download_name', f'document.{file_type}')
    output_ext = download_name.rsplit('.', 1)[-1] if '.' in download_name else file_type
    output_path = os.path.join(_get_output_dir(), f'{task_id}.{output_ext}')
    if not os.path.exists(output_path):
        raise Http404('Converted file not found.')

    try:
        with open(output_path, 'rb') as f:
            data = f.read()
    except Exception:
        raise Http404('Could not read converted file.')

    mime = task.get('mime', 'application/octet-stream')

    response = HttpResponse(data, content_type=mime)
    response['Content-Disposition'] = f'attachment; filename="{download_name}"'
    response['Content-Length'] = len(data)

    cache.delete(cache_key)
    _cleanup_task(task_id)
    return response
