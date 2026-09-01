import os
import io
import zipfile
import json
from unittest.mock import patch
from django.test import TestCase, Client
from django.conf import settings
from django.core.files.uploadedfile import SimpleUploadedFile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from PIL import Image as PilImage
from docx import Document


def _create_pdf_with_text(text_lines, page_size=letter):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=page_size)
    width, height = page_size
    y = height - 72
    for line in text_lines:
        c.drawString(72, y, line)
        y -= 14
    c.save()
    buffer.seek(0)
    return buffer


def _create_pdf_with_image(text_lines, image_bytes, page_size=letter):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=page_size)
    width, height = page_size
    y = height - 72
    for line in text_lines:
        c.drawString(72, y, line)
        y -= 14
    img = PilImage.open(io.BytesIO(image_bytes))
    img_w, img_h = img.size
    max_w = width - 144
    max_h = 200
    ratio = min(max_w / img_w, max_h / img_h, 1)
    c.drawImage(ImageReader(io.BytesIO(image_bytes)), 72, y - max_h * ratio, width=img_w * ratio, height=img_h * ratio, preserveAspectRatio=True, mask='auto')
    c.save()
    buffer.seek(0)
    return buffer


def _create_dummy_image():
    img_buffer = io.BytesIO()
    img = PilImage.new('RGB', (200, 100), color='red')
    img.save(img_buffer, format='PNG')
    img_buffer.seek(0)
    return img_buffer.getvalue()


def _create_docx_with_text(text_lines):
    doc = Document()
    doc.add_heading('Test Document', 0)
    for line in text_lines:
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


class DocConverterIndexTests(TestCase):
    def setUp(self):
        self.client = Client()

    def test_index_page_loads(self):
        response = self.client.get('/tools/pdf-to-word/')
        self.assertEqual(response.status_code, 200)
        self.assertIn(b'Document Converter', response.content)


class DocConverterUploadTests(TestCase):
    def setUp(self):
        self.client = Client()
        self.url = '/tools/pdf-to-word/upload/'

    def test_non_pdf_rejected(self):
        uploaded = SimpleUploadedFile('test.txt', b'hello world', content_type='text/plain')
        response = self.client.post(self.url, {'file': uploaded})
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertFalse(data['success'])
        self.assertIn('PDF and Word', data['error'])

    def test_oversized_file_rejected(self):
        big_file = SimpleUploadedFile('huge.pdf', b'0' * (20 * 1024 * 1024 + 1), content_type='application/pdf')
        response = self.client.post(self.url, {'file': big_file})
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertFalse(data['success'])
        self.assertIn('too large', data['error'].lower())

    def test_corrupted_pdf_handled(self):
        bad_file = SimpleUploadedFile('fake.pdf', b'This is not a PDF file at all.', content_type='application/pdf')
        response = self.client.post(self.url, {'file': bad_file})
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertFalse(data['success'])
        self.assertIn('corrupted', data['error'].lower())

    def test_successful_pdf_to_docx_conversion(self):
        pdf_buffer = _create_pdf_with_text(['Hello World', 'This is a test.'])
        uploaded = SimpleUploadedFile('test.pdf', pdf_buffer.read(), content_type='application/pdf')
        response = self.client.post(self.url, {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['success'])
        self.assertEqual(data['file_type'], 'pdf')
        self.assertRegex(data['download_url'], r'/tools/pdf-to-word/download/[^/]+/')

        download_response = self.client.get(data['download_url'])
        self.assertEqual(download_response.status_code, 200)
        self.assertIn('wordprocessingml', download_response['Content-Type'])
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(download_response.content)))

    def test_successful_docx_to_pdf_conversion(self):
        docx_buffer = _create_docx_with_text(['Hello from Word', 'This is a test.'])
        uploaded = SimpleUploadedFile('test.docx', docx_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        with patch('doc_converter.views._convert_docx_to_pdf') as mock_convert:
            mock_convert.side_effect = lambda src, dst: open(dst, 'wb').write(b'%PDF-1.4\n')
            response = self.client.post(self.url, {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['success'])
        self.assertEqual(data['file_type'], 'docx')
        self.assertRegex(data['download_url'], r'/tools/pdf-to-word/download/[^/]+/')

        download_response = self.client.get(data['download_url'])
        self.assertEqual(download_response.status_code, 200)
        self.assertEqual(download_response['Content-Type'], 'application/pdf')
        self.assertTrue(len(download_response.content) > 0)

    def test_successful_conversion_with_images(self):
        image_bytes = _create_dummy_image()
        pdf_buffer = _create_pdf_with_image(['Hello World', 'This is a test.'], image_bytes)
        uploaded = SimpleUploadedFile('test_image.pdf', pdf_buffer.read(), content_type='application/pdf')
        response = self.client.post(self.url, {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['success'])
        self.assertIn('download_url', data)

        download_response = self.client.get(data['download_url'])
        self.assertEqual(download_response.status_code, 200)
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(download_response.content)))
        with zipfile.ZipFile(io.BytesIO(download_response.content)) as zf:
            self.assertIn('word/document.xml', zf.namelist())

    def test_download_cleans_up_and_cannot_reuse(self):
        pdf_buffer = _create_pdf_with_text(['Test'])
        uploaded = SimpleUploadedFile('cleanup_test.pdf', pdf_buffer.read(), content_type='application/pdf')
        response = self.client.post(self.url, {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['success'])

        download_response = self.client.get(data['download_url'])
        self.assertEqual(download_response.status_code, 200)

        second_download = self.client.get(data['download_url'])
        self.assertEqual(second_download.status_code, 404)

    def test_expired_task_returns_404(self):
        response = self.client.get('/tools/pdf-to-word/download/00000000-0000-0000-0000-000000000000/')
        self.assertEqual(response.status_code, 404)


class DocConverterEditTests(TestCase):
    def setUp(self):
        self.client = Client()

    def test_edit_docx_returns_text(self):
        docx_buffer = _create_docx_with_text(['First paragraph', 'Second paragraph', 'Third paragraph'])
        uploaded = SimpleUploadedFile('edit_test.docx', docx_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response = self.client.post('/tools/pdf-to-word/upload/', {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['success'])
        task_id = data['task_id']

        edit_response = self.client.post(f'/tools/pdf-to-word/edit/{task_id}/')
        self.assertEqual(edit_response.status_code, 200)
        edit_data = edit_response.json()
        self.assertTrue(edit_data['success'])
        self.assertIn('First paragraph', edit_data['text'])
        self.assertIn('Second paragraph', edit_data['text'])
        self.assertIn('Third paragraph', edit_data['text'])

    def test_edit_pdf_task_returns_404(self):
        pdf_buffer = _create_pdf_with_text(['PDF text'])
        uploaded = SimpleUploadedFile('edit_pdf.pdf', pdf_buffer.read(), content_type='application/pdf')
        response = self.client.post('/tools/pdf-to-word/upload/', {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        task_id = data['task_id']

        edit_response = self.client.post(f'/tools/pdf-to-word/edit/{task_id}/')
        self.assertEqual(edit_response.status_code, 404)

    def test_save_edited_docx_returns_download(self):
        docx_buffer = _create_docx_with_text(['Original text'])
        uploaded = SimpleUploadedFile('save_edit.docx', docx_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response = self.client.post('/tools/pdf-to-word/upload/', {'file': uploaded})
        self.assertEqual(response.status_code, 200)
        data = response.json()
        task_id = data['task_id']

        save_response = self.client.post(
            '/tools/pdf-to-word/save-edited/',
            data=json.dumps({'task_id': task_id, 'text': '[H1] New Title\n\nNew paragraph text.'}),
            content_type='application/json',
        )
        self.assertEqual(save_response.status_code, 200)
        save_data = save_response.json()
        self.assertTrue(save_data['success'])
        self.assertIn('download_url', save_data)

        dl = self.client.get(save_data['download_url'])
        self.assertEqual(dl.status_code, 200)
        self.assertIn('wordprocessingml', dl['Content-Type'])
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(dl.content)))


class DocConverterBrowserFlowTests(TestCase):
    def setUp(self):
        self.client = Client()

    def test_page_contains_ajax_form_and_js(self):
        response = self.client.get('/tools/pdf-to-word/')
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('id="convertForm"', content)
        self.assertIn('id="fileInput"', content)
        self.assertIn('preventDefault', content)
        self.assertIn('XMLHttpRequest', content)
        self.assertIn('id="progressWrapper"', content)
        self.assertIn('id="resultCard"', content)

    def test_form_not_rendered_in_head(self):
        response = self.client.get('/tools/pdf-to-word/')
        content = response.content.decode()
        head_end = content.find('</head>')
        body_start = content.find('<body')
        form_id_pos = content.find('id="convertForm"')
        self.assertGreater(head_end, 0, 'Page should have </head>')
        self.assertGreater(body_start, 0, 'Page should have <body')
        self.assertGreater(form_id_pos, 0, 'Page should have convertForm')
        self.assertGreater(form_id_pos, head_end, 'Form should be inside <body>, not inside <head>')

    def test_full_browser_flow_success(self):
        pdf_buffer = _create_pdf_with_text(['Browser flow test'])
        uploaded = SimpleUploadedFile('flow.pdf', pdf_buffer.read(), content_type='application/pdf')

        get_response = self.client.get('/tools/pdf-to-word/')
        self.assertEqual(get_response.status_code, 200)

        post_response = self.client.post('/tools/pdf-to-word/upload/', {'file': uploaded})
        self.assertEqual(post_response.status_code, 200)
        data = post_response.json()
        self.assertTrue(data['success'])
        self.assertIn('download_url', data)

        dl = self.client.get(data['download_url'])
        self.assertEqual(dl.status_code, 200)
        self.assertIn('application/vnd.openxmlformats-officedocument.wordprocessingml.document', dl['Content-Type'])
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(dl.content)))

    def test_full_browser_flow_error_shows_message(self):
        bad_file = SimpleUploadedFile('bad.pdf', b'not a pdf', content_type='application/pdf')
        response = self.client.post('/tools/pdf-to-word/upload/', {'file': bad_file})
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertFalse(data['success'])
        self.assertIn('error', data)